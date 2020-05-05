from docx import Document

def replaceFromDocx(file_name_to_be_changed=str, file_name_to_be_changed_into=str, replaceTo=[], replaceWith=[]):
    doc = Document(file_name_to_be_changed)
    new_doc = Document()

    for i in doc.paragraphs:
        #print(i.text)
        para = new_doc.add_paragraph()
        text = i.text
        for word_to_be_replaced in replaceTo:
            ...


        para.add_run(text)

    new_doc.save('new.docx')
