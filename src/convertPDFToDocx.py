import PyPDF2
from docx import Document

def PdfDocxcoverter(pdfName):
    docName = pdfName.split(".")[0] + ".docx"
    fopen = open(pdfName, 'rb')
    pdfReader = PyPDF2.PdfFileReader(fopen)

    doc = Document()

    for i in range(pdfReader.numPages):
        page = pdfReader.getPage(i).extractText()
        print(page)
        paragraph = doc.add_paragraph()
        paragraph.add_run(page)

    doc.save(docName)


